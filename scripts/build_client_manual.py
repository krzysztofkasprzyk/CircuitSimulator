from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "assets"
OUTPUT = ROOT / "docs" / "Instrukcja_RelayLab.docx"

INK = "20242A"
RED = "E3222E"
BLUE = "1769AA"
GREEN = "15803D"
AMBER = "9A5B00"
LIGHT = "F3F5F7"
MUTED = RGBColor(98, 106, 116)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def font(run, size=None, bold=None, color=None, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color) if isinstance(color, str) else color
    return run


def page_title(doc, number, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    badge = p.add_run(f" {number} ")
    font(badge, 11, True, "FFFFFF")
    badge._element.get_or_add_rPr().append(OxmlElement("w:highlight"))
    badge._element.rPr[-1].set(qn("w:val"), "red")
    font(p.add_run(f"  {title}"), 20, True, INK)
    q = doc.add_paragraph(subtitle)
    q.paragraph_format.space_after = Pt(7)
    font(q.runs[0], 9.5, False, MUTED)


def add_image(doc, filename, alt_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    shape = run.add_picture(str(ASSETS / filename), width=Inches(8.75))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text)


def steps_table(doc, steps):
    table = doc.add_table(rows=1, cols=len(steps))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    colors = [RED, BLUE, GREEN, AMBER]
    total = Inches(9.9)
    for index, text in enumerate(steps):
        cell = table.rows[0].cells[index]
        cell.width = total / len(steps)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, 90, 120, 90, 120)
        shade(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(f"{index + 1}. ")
        font(r, 10, True, colors[index])
        font(p.add_run(text), 9.2, False, INK)
    return table


def note(doc, title, body, color=RED):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.45)
    table.columns[1].width = Inches(8.4)
    left, right = table.rows[0].cells
    shade(left, color)
    shade(right, "FFF5F6" if color == RED else LIGHT)
    for cell in (left, right):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, 105, 130, 105, 130)
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(lp.add_run(title), 10, True, "FFFFFF")
    rp = right.paragraphs[0]
    font(rp.add_run(body), 9.5, False, INK)


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.top_margin = Cm(1.25)
section.bottom_margin = Cm(1.15)
section.left_margin = Cm(1.4)
section.right_margin = Cm(1.4)
section.header_distance = Cm(0.45)
section.footer_distance = Cm(0.45)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
styles["Normal"].font.size = Pt(10)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(hp.add_run("RELAYLAB  |  INSTRUKCJA UŻYTKOWNIKA"), 8, True, MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(fp.add_run("RelayLab v2  •  Symulator obwodów przekaźnikowych  •  lipiec 2026"), 8, False, MUTED)

# Okładka
band = doc.add_table(rows=1, cols=1)
band.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = band.cell(0, 0)
shade(cell, INK)
set_cell_margins(cell, 350, 380, 350, 380)
p = cell.paragraphs[0]
font(p.add_run("RelayLab"), 32, True, "FFFFFF")
p2 = cell.add_paragraph()
font(p2.add_run("Projektowanie i symulacja obwodów przekaźnikowych"), 15, False, "FFFFFF")
p3 = cell.add_paragraph()
p3.paragraph_format.space_before = Pt(10)
font(p3.add_run("PROSTA INSTRUKCJA DLA UŻYTKOWNIKA"), 10, True, "F3A1A7")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
font(p.add_run("Wersja online"), 11, True, RED)
u = doc.add_paragraph("https://krzysztofkasprzyk.github.io/CircuitSimulator/")
u.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(u.runs[0], 13, True, BLUE)

doc.add_paragraph()
steps_table(doc, [
    "Otwórz adres programu w aktualnej przeglądarce Edge, Chrome lub Firefox.",
    "Zbuduj schemat w trybie Projektowanie, korzystając z biblioteki po lewej.",
    "Włącz Symulację i obsługuj przyciski z pulpitu po prawej.",
    "Zapisuj kopię projektu przyciskiem Eksport JSON."
])
doc.add_paragraph()
note(doc, "WAŻNE", "Program symuluje logikę obwodów przekaźnikowych. Nie zastępuje obliczeń zwarciowych, doboru zabezpieczeń ani projektu wykonawczego.")

# Ekran główny
doc.add_page_break()
page_title(doc, "1", "Poznaj ekran programu", "Wszystkie najważniejsze funkcje są stale widoczne - nie trzeba otwierać dodatkowych okien.")
add_image(doc, "01-interfejs.png", "Główny ekran RelayLab z zaznaczoną biblioteką, zakładkami, arkuszem i panelem właściwości")
steps_table(doc, [
    "Biblioteka: wyszukaj i dodaj element.",
    "Zakładki: przełączaj, dodawaj i nazywaj arkusze.",
    "Arkusz: układaj elementy i rysuj przewody.",
    "Panel: steruj układem i zmieniaj właściwości."
])

# Edycja
doc.add_page_break()
page_title(doc, "2", "Dodaj element i narysuj przewód", "Element trafia w pierwsze wolne miejsce. Wszystkie obiekty są automatycznie wyrównywane do siatki.")
add_image(doc, "02-edycja.png", "Dodawanie, przesuwanie i łączenie elementów na arkuszu RelayLab")
steps_table(doc, [
    "Kliknij element w bibliotece.",
    "Przeciągnij go w wybrane miejsce.",
    "Kliknij zacisk pierwszego elementu, potem drugiego.",
    "Ustaw nazwę, typ i parametry po prawej."
])

# Właściwości
doc.add_page_break()
page_title(doc, "3", "Powiąż cewkę z zestykami", "Oznaczenie funkcjonalne jest kluczem sterowania - identyczna nazwa oznacza wspólny przekaźnik.")
add_image(doc, "03-wlasciwosci.png", "Cewka K1 i panel jej właściwości na drugim arkuszu")
steps_table(doc, [
    "Wybierz właściwy arkusz.",
    "Kliknij cewkę lub zestyk.",
    "Nadaj wszystkim elementom wspólne oznaczenie, np. K1.",
    "Dla cewek czasowych ustaw zwłokę w milisekundach."
])

# Symulacja
doc.add_page_break()
page_title(doc, "4", "Uruchom i obsługuj symulację", "Czerwony kolor pokazuje zamkniętą, aktywną drogę prądu. Czarny oznacza brak zasilenia.")
add_image(doc, "04-symulacja.png", "Aktywna symulacja RelayLab z czerwonymi przewodami i przyciskiem START")
steps_table(doc, [
    "Kliknij Symulacja na górnym pasku.",
    "Użyj przycisków z pulpitu sterowania.",
    "Obserwuj czerwone przewody i zmianę zestyków.",
    "Reset stanów zeruje przyciski, remanencję i pamięć impulsową."
])

# Zapis i pomoc
doc.add_page_break()
page_title(doc, "5", "Zapis, wymiana projektu i szybka pomoc", "Projekt jest zapisywany automatycznie w tej przeglądarce. Eksport JSON tworzy przenośną kopię.")
steps_table(doc, [
    "Zapisz utrwala bieżący projekt lokalnie.",
    "Eksport JSON pobiera plik projektu na komputer.",
    "Import JSON otwiera projekt otrzymany od innej osoby.",
    "Projekt przykładowy odtwarza gotowy obwód szkoleniowy."
])
doc.add_paragraph()
note(doc, "GDY NIE DZIAŁA", "Sprawdź, czy każdy tor kończy się powrotem 0 V, czy zaciski są połączone przewodami oraz czy cewka i jej zestyki mają dokładnie tę samą nazwę.", BLUE)
doc.add_paragraph()
note(doc, "LAMPY", "Opcja „Symuluj przepalenie” tworzy przerwę w obwodzie. Żarówka sterowana cewką kontroli światła musi mieć takie samo oznaczenie jak ta cewka.", GREEN)
doc.add_paragraph()
note(doc, "ARKUSZE", "Łączniki międzyarkuszowe oznaczone tą samą literą są elektrycznie połączone, nawet gdy znajdują się na różnych zakładkach.", AMBER)
doc.add_paragraph()
note(doc, "KOPIA DLA KLIENTA", "Przed większą zmianą użyj Eksport JSON. Plik można później przywrócić przez Import JSON bez przebudowywania schematu.")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "RelayLab - instrukcja użytkownika"
doc.core_properties.subject = "Obsługa symulatora obwodów przekaźnikowych RelayLab v2"
doc.core_properties.author = "RelayLab"
doc.core_properties.keywords = "RelayLab, przekaźniki, symulator, instrukcja"
doc.save(OUTPUT)
print(f"Zapisano {OUTPUT}")
